import os
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
import PIL
from PIL import Image
import validators
from pathvalidate import sanitize_filepath
from docx.oxml.ns import qn

#----------------------------------------------------------

def strip_url(URL):
    global ln_name, main_rawparsed

    # Get main URL and main ln name
    main_URL = re.sub('\/c.*$','',URL)
    main_rawconts = requests.get(main_URL)
    main_rawparsed = BeautifulSoup(main_rawconts.content,'html.parser')
    ln_name = main_rawparsed.find("span",attrs='series-name').text.strip()

def find_select_volume():
    global vol_URL, baseURL

    baseURL = 'https://ln.hako.vn'
    volnum = 1

    # Find all link in page
    vol_links = main_rawparsed.find_all("a",href=True)
    
    # Set dictionaries for saving URL and name
    vol_URLdict = {}
    vol_namedict = {}
    print("Get the volume:")

    # Find volumes
    for x in vol_links:

        # Find the volumes' link
        if (re.search('/t[0-9]',x['href']))!= None:
            # Set vol's link
            vol_URL = baseURL + x['href'] 
            vol_rawconts = requests.get(vol_URL)
            vol_rawparsed = BeautifulSoup(vol_rawconts.content,'html.parser')

            # Find vol's name
            vol_name = vol_rawparsed.find("span",attrs='volume-name').text.strip()
            print(f'{volnum}. {vol_name}')

            # Save to dicts for future uses
            vol_URLdict[volnum] = vol_URL
            vol_namedict[volnum] = vol_name
            volnum +=1

    # Select volume to get
    vol_selection = ""
    while vol_selection not in range(volnum):
        try:
            print("Choose volume:\n>", end=" ")
            vol_selection = int(input())
        except ValueError:
            print("Invalid volume. Please enter again:")
            pass
    print(f"Chosen the volume: {vol_namedict[vol_selection]}")
    vol_URL = vol_URLdict[vol_selection]
    print(f"URL: {vol_URL}")

#----------------------------------------------------------

def find_URL_chap():
    global chap_URL

    chap_rawconts = requests.get(vol_URL)
    chap_rawparsed = BeautifulSoup(chap_rawconts.content,'html.parser')
    chap_URL_raw = chap_rawparsed.find_all("a",href=True)
    chap_URL = []
    for x in chap_URL_raw:
        if (re.search('/c[0-9]',x['href']))!= None:
            chap_URL.append(baseURL+x['href'])
    
#----------------------------------------------------------

def get_contents(URL):
    print("Loading...")
    rawconts = requests.get(URL)
    rawparsed = BeautifulSoup(rawconts.content, "html.parser")
    
    # Pull the content from web
    vol_name = rawparsed.find("h2",attrs="title-item").contents[0].text
    chap_name = rawparsed.find("h4",attrs="title-item").contents[0].text
    ln_conts = rawparsed.find_all("p",id=True)
    notes = rawparsed.find_all(attrs="note-content long-text")
    

    # Create an instance of Document
    document = Document()

    # Edit styles
    styles = document.styles
    styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
    styles['Heading 1'].font.size = Pt(14)
    styles['Heading 1'].font.name = "Cambria"
    styles['Normal'].font.name = "Cambria"
    styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    styles['Normal'].font.size = Pt(12)

    # Add heading + change font to Cambria
    head = document.add_heading(chap_name,1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style = head.style
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Cambria")
    rFonts.set(qn("w:hAnsiTheme"), "Cambria")

    # Default value 
    img_num = 1  
    note_num = 1
    
    # Make the default folder
    d = os.path.dirname(__file__) # directory of script
    pathname = f'{d}/{ln_name}/{vol_name}/{sanitize_filepath(chap_name)}'   
    imgfolder = f'{pathname}/image'
    os.makedirs(os.path.dirname(imgfolder),exist_ok=True)

    # Add contents
    for cnt in ln_conts:
        for cont in cnt.contents:
            # Make the image file
            imgname = f'{imgfolder}/img{img_num}.jpg'
            os.makedirs(os.path.dirname(imgname), exist_ok=True)
            
            # Find the image
            if cont.name == 'img':
                img_url = cont['src']   
                img_data = requests.get(img_url, headers={'referer': "https://ln.hako.vn"}).content
                
                # Get the image
                try:
                    with open(imgname,'wb+') as img:
                        img.write(img_data)
                    img.close()
                    im = PIL.Image.open(imgname)
                    if im.mode in ("RGBA", "P"): 
                        im = im.convert("RGB")
                    im.save(imgname)
                
                # Delete the image file if can't get the file
                except PIL.UnidentifiedImageError:
                    document.add_paragraph(f"img{img_num}.jpg (get the image by hand)")
                    os.remove(imgname)

                # Save the image
                else:
                    document.add_paragraph(f"img{img_num}.jpg (on the folder)")
                    img_num += 1

        # Add paragraph for text + some alignment
        parrg = document.add_paragraph(style='Normal')
        parrg.paragraph_format.line_spacing = 1.0
        parrg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY        
        
        # Add text in paragraph word by word
        for word in cnt:
            run = parrg.add_run() # Allow making word Bold, Italic, Strikethrough
            strip_text=word.text.replace(u'\u200e','')
            
            # Search for [note...] and replace
            test = re.search('\[note[0-9][0-9][0-9][0-9][0-9]\]',word.text) # Regex for [note...]
            if test!= None:
                k = word.text.replace(test.group(0),f" (Note {note_num})")
                note_num+=1
                run.add_text(k)
            else:
                run.add_text(strip_text)
            
            # Making word Bold, Italic, Strikethrough
            if word.name == 's':
                run.font.strike = True
                continue
            elif word.name == 'strong':
                run.bold = True
                continue
            elif word.name == 'em':
                run.italic = True
                continue
    
    # Add notes at the end of document
    if notes:
        document.add_paragraph("\nNote:")
        stt=1
        for note in notes:
            nte = note.text.strip()
            document.add_paragraph(f"{stt}. {nte}",style='Normal')
            stt+=1

    # Remove the 'image' folder if there is no image            
    if not os.listdir(imgfolder):
        os.rmdir(imgfolder)
    
    # Save the document        
    document.save(f"{pathname}/{sanitize_filepath(chap_name)}.docx")

#----------------------------------------------------------

def user_interface():
    inp = ""
    while (inp not in ["1","2","3"]):
        print("1: Get a whole volume\n2: Get a chapter\n3: Exit\n>",end=" ")
        inp = input()
        if (inp not in ["1","2","3"]):
            print("Invalid. Please enter your choice: ")
            continue
        elif inp=="1":
            print("Please enter LN's URL:")
            print(">",end=" ")
            URL = input()
            while not validators.url(URL):
                print("Invalid URL. Please enter URL:")
                URL = input()
            else:
                strip_url(URL)
                find_select_volume()
                find_URL_chap()
                for x in chap_URL:
                    get_contents(x)
                print("Done! Please enter your choice:")
                inp = ""
        elif inp=="2":
            print("Please enter your URL:")
            print(">",end=" ")
            URL = input()
            while not validators.url(URL):
                print("Invalid URL. Please enter URL:")
                URL = input()
            else:
                strip_url(URL)
                get_contents(URL)
                print("Done! Please enter your choice:")
                inp = ""
        elif inp=="3":
            break

#---------------------------------------------------------------------------------------
# Callout function
print(" _  _   _   _  _____    _         ___   ___   _____  __")
print("| || | /_\ | |/ / _ \  | |_ ___  |   \ / _ \ / __\ \/ /")
print("| __ |/ _ \| ' < (_) | |  _/ _ \ | |) | (_) | (__ >  < ")
print("|_||_/_/ \_\_|\_\___/   \__\___/ |___/ \___/ \___/_/\_\ ")
print()    
print("Please enter your choice:")
user_interface()
